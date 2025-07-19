import PyPDF2
import docx
import chardet
import aiofiles
from pathlib import Path
from typing import List, Dict, Any
import logging
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
        self.max_file_size = 50 * 1024 * 1024  # 50MB
        
    async def process_pdf(self, file_path: Path) -> str:
        """
        Extrai texto de PDF usando PyPDF2
        """
        try:
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n"
                    
            return self._clean_extracted_text(text_content)
            
        except Exception as e:
            logger.error(f"Erro ao processar PDF {file_path}: {str(e)}")
            raise Exception(f"Erro ao processar PDF: {str(e)}")
    
    async def process_docx(self, file_path: Path) -> str:
        """
        Extrai texto de DOCX preservando formatação básica
        """
        try:
            doc = docx.Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
                    
            # Processar tabelas também
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"
                    
            return self._clean_extracted_text(text_content)
            
        except Exception as e:
            logger.error(f"Erro ao processar DOCX {file_path}: {str(e)}")
            raise Exception(f"Erro ao processar DOCX: {str(e)}")
    
    async def process_txt(self, file_path: Path) -> str:
        """
        Processa TXT com detecção automática de encoding
        """
        try:
            # Detectar encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result['encoding'] or 'utf-8'
            
            # Ler arquivo com encoding detectado
            async with aiofiles.open(file_path, 'r', encoding=encoding) as file:
                text_content = await file.read()
                
            return self._clean_extracted_text(text_content)
            
        except Exception as e:
            logger.error(f"Erro ao processar TXT {file_path}: {str(e)}")
            raise Exception(f"Erro ao processar TXT: {str(e)}")
    
    def _clean_extracted_text(self, text: str) -> str:
        """
        Limpa e normaliza o texto extraído
        """
        # Remover múltiplas quebras de linha
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # Remover espaços excessivos
        text = re.sub(r' +', ' ', text)
        
        # Remover caracteres especiais problemáticos
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)
        
        return text.strip()
    
    def intelligent_chunking(self, text: str, metadata: Dict[str, Any], chunk_size: int = 1000) -> List[Dict[str, Any]]:
        """
        Chunking inteligente respeitando estrutura legal
        """
        chunks = []
        
        # Dividir por artigos/seções se possível
        if self._has_legal_structure(text):
            chunks = self._chunk_by_legal_structure(text, metadata, chunk_size)
        else:
            chunks = self._chunk_by_paragraphs(text, metadata, chunk_size)
            
        return chunks
    
    def _has_legal_structure(self, text: str) -> bool:
        """
        Verifica se o texto tem estrutura legal típica
        """
        patterns = [
            r'Artigo\s+\d+',
            r'Art\.\s+\d+',
            r'Capítulo\s+[IVX]+',
            r'Secção\s+[IVX]+',
            r'§\s*\d+',
            r'^\d+\.\s+',  # Numeração simples
        ]
        
        for pattern in patterns:
            if re.search(pattern, text, re.IGNORECASE | re.MULTILINE):
                return True
        return False
    
    def _chunk_by_legal_structure(self, text: str, metadata: Dict[str, Any], max_size: int) -> List[Dict[str, Any]]:
        """
        Divide texto respeitando estrutura de artigos/seções
        """
        chunks = []
        
        # Padrões para dividir
        patterns = [
            r'(Artigo\s+\d+[^\n]*)',
            r'(Art\.\s+\d+[^\n]*)',
            r'(Capítulo\s+[IVX]+[^\n]*)',
            r'(Secção\s+[IVX]+[^\n]*)',
        ]
        
        sections = []
        current_section = ""
        
        lines = text.split('\n')
        
        for line in lines:
            is_new_section = False
            
            for pattern in patterns:
                if re.match(pattern, line.strip(), re.IGNORECASE):
                    if current_section.strip():
                        sections.append(current_section.strip())
                    current_section = line + '\n'
                    is_new_section = True
                    break
            
            if not is_new_section:
                current_section += line + '\n'
        
        # Adicionar última seção
        if current_section.strip():
            sections.append(current_section.strip())
        
        # Se não encontrou seções, usar divisão por parágrafos
        if len(sections) <= 1:
            return self._chunk_by_paragraphs(text, metadata, max_size)
        
        # Criar chunks das seções
        for i, section in enumerate(sections):
            if len(section) <= max_size:
                chunks.append({
                    'text': section,
                    'chunk_index': i,
                    'metadata': {
                        **metadata,
                        'chunk_type': 'legal_section',
                        'section_number': i + 1
                    }
                })
            else:
                # Seção muito grande, dividir mais
                sub_chunks = self._split_large_section(section, metadata, max_size, i)
                chunks.extend(sub_chunks)
        
        return chunks
    
    def _chunk_by_paragraphs(self, text: str, metadata: Dict[str, Any], max_size: int) -> List[Dict[str, Any]]:
        """
        Divide texto por parágrafos quando não há estrutura legal clara
        """
        chunks = []
        paragraphs = text.split('\n\n')
        
        current_chunk = ""
        chunk_index = 0
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            # Se adicionar este parágrafo exceder o limite
            if len(current_chunk) + len(paragraph) + 2 > max_size:
                if current_chunk:
                    chunks.append({
                        'text': current_chunk.strip(),
                        'chunk_index': chunk_index,
                        'metadata': {
                            **metadata,
                            'chunk_type': 'paragraph',
                            'chunk_number': chunk_index + 1
                        }
                    })
                    chunk_index += 1
                
                # Se o parágrafo é muito grande sozinho
                if len(paragraph) > max_size:
                    sub_chunks = self._split_large_paragraph(paragraph, metadata, max_size, chunk_index)
                    chunks.extend(sub_chunks)
                    chunk_index += len(sub_chunks)
                    current_chunk = ""
                else:
                    current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += '\n\n' + paragraph
                else:
                    current_chunk = paragraph
        
        # Adicionar último chunk
        if current_chunk:
            chunks.append({
                'text': current_chunk.strip(),
                'chunk_index': chunk_index,
                'metadata': {
                    **metadata,
                    'chunk_type': 'paragraph',
                    'chunk_number': chunk_index + 1
                }
            })
        
        return chunks
    
    def _split_large_section(self, section: str, metadata: Dict[str, Any], max_size: int, base_index: int) -> List[Dict[str, Any]]:
        """
        Divide seção legal muito grande em chunks menores
        """
        chunks = []
        sentences = section.split('. ')
        
        current_chunk = ""
        sub_index = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            if not sentence.endswith('.'):
                sentence += '.'
                
            if len(current_chunk) + len(sentence) + 2 > max_size:
                if current_chunk:
                    chunks.append({
                        'text': current_chunk.strip(),
                        'chunk_index': f"{base_index}.{sub_index}",
                        'metadata': {
                            **metadata,
                            'chunk_type': 'legal_subsection',
                            'parent_section': base_index + 1,
                            'subsection_number': sub_index + 1
                        }
                    })
                    sub_index += 1
                current_chunk = sentence
            else:
                if current_chunk:
                    current_chunk += ' ' + sentence
                else:
                    current_chunk = sentence
        
        if current_chunk:
            chunks.append({
                'text': current_chunk.strip(),
                'chunk_index': f"{base_index}.{sub_index}",
                'metadata': {
                    **metadata,
                    'chunk_type': 'legal_subsection',
                    'parent_section': base_index + 1,
                    'subsection_number': sub_index + 1
                }
            })
        
        return chunks
    
    def _split_large_paragraph(self, paragraph: str, metadata: Dict[str, Any], max_size: int, base_index: int) -> List[Dict[str, Any]]:
        """
        Divide parágrafo muito grande
        """
        chunks = []
        words = paragraph.split()
        
        current_chunk = ""
        sub_index = 0
        
        for word in words:
            if len(current_chunk) + len(word) + 1 > max_size:
                if current_chunk:
                    chunks.append({
                        'text': current_chunk.strip(),
                        'chunk_index': f"{base_index}.{sub_index}",
                        'metadata': {
                            **metadata,
                            'chunk_type': 'paragraph_fragment',
                            'fragment_number': sub_index + 1
                        }
                    })
                    sub_index += 1
                current_chunk = word
            else:
                if current_chunk:
                    current_chunk += ' ' + word
                else:
                    current_chunk = word
        
        if current_chunk:
            chunks.append({
                'text': current_chunk.strip(),
                'chunk_index': f"{base_index}.{sub_index}",
                'metadata': {
                    **metadata,
                    'chunk_type': 'paragraph_fragment', 
                    'fragment_number': sub_index + 1
                }
            })
        
        return chunks